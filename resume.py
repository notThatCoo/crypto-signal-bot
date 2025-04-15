from docx import Document
from docx.shared import Inches

# Create a new Document
doc = Document()

# Header
doc.add_heading("Tolulope “Tolu” Adeleye", 0)
doc.add_paragraph("Inglewood, CA | (424) 223-0130 | tadel002@ucr.edu")
doc.add_paragraph("LinkedIn: https://www.linkedin.com/in/tolulope-adeleye-1600321b4 | GitHub: https://github.com/notThatCoo")

# Education
doc.add_heading("Education", level=1)
doc.add_paragraph("BS in Data Science, University of California, Riverside — Expected June 2026")
doc.add_paragraph("Relevant Coursework: Optimization, Software Construction, Machine Learning, Data Structures & Algorithms")
doc.add_paragraph("Affiliations: National Society of Black Engineers (NSBE)")

# Technical Skills
doc.add_heading("Technical Skills", level=1)
doc.add_paragraph(
    "Languages & Tools: Python, C/C++, SQL, MATLAB, Git, Bash, Jupyter, Google Colab\n"
    "Libraries: Pandas, NumPy, scikit-learn, Keras, matplotlib, ccxt\n"
    "Concepts: Machine Learning, Statistical Inference, Optimization, Cryptographic Algorithms, Time Series, Backtesting\n"
    "Platforms: GitHub, Google Cloud VMs, Discord Bots, Kraken API, Linux (SSH/cron)"
)

# Projects
doc.add_heading("Projects", level=1)
doc.add_paragraph("Live Crypto Trading Bot with Backtesting and ML", style='List Bullet')
doc.add_paragraph(
    "Jan 2024 – Present | GitHub: https://github.com/notThatCoo/crypto-signal-bot\n"
    "- Designed and deployed a live crypto trading bot using a Random Forest model trained on volatility and momentum indicators\n"
    "- Integrated Discord for hourly predictions, logs, and wallet updates using a Google Cloud VM\n"
    "- Simulated live capital flow using a custom wallet class with full trade logging, fees, and margin shorting via Kraken API\n"
    "- Implemented prediction logging, model performance summaries (accuracy, Sharpe ratio), and backtesting infrastructure\n"
    "- Planned enhancements include optimizer integration, SQL-based trade storage, and deployment scalability"
)

doc.add_paragraph("RPG Game Engine (C++)", style='List Bullet')
doc.add_paragraph(
    "Feb 2024\n"
    "- Collaborated in a team to implement a C++ role-playing game engine using classes, inheritance, and state management\n"
    "- Led as Scrum master by managing task delegation, GitHub commits, and sprint communication"
)

doc.add_paragraph("RSA Encryption Tool (C++)", style='List Bullet')
doc.add_paragraph(
    "Jan 2024\n"
    "- Developed a command-line RSA encryption tool and optimized runtime via modular exponentiation\n"
    "- Gained practical experience in algorithm design, number theory, and basic cybersecurity"
)

doc.add_paragraph("CS50 Certificate (Harvardx/edX)", style='List Bullet')
doc.add_paragraph(
    "Summer 2022\n"
    "- Completed foundational programming in C, Python, and SQL; covered data structures, memory management, and web development"
)

# Leadership
doc.add_heading("Leadership & Involvement", level=1)
doc.add_paragraph("NSBE – National Society of Black Engineers", style='List Bullet')
doc.add_paragraph(
    "UCR Chapter\n"
    "- Serve on outreach and event planning committees to support inclusion and growth in STEM\n"
    "- Mentor new members and lead educational workshops in Python, ML, and algorithm fundamentals"
)

doc.add_paragraph("Hackathon: Rosehack", style='List Bullet')
doc.add_paragraph(
    "- Developed data analysis pipelines using Pandas in a time-constrained hackathon setting\n"
    "- Strengthened rapid prototyping skills with Jupyter and Google Colab"
)

# Objective
doc.add_heading("Objective", level=1)
doc.add_paragraph(
    "I’m seeking a machine learning, trading systems, or data-driven software engineering internship where I can "
    "combine applied statistics, optimization, and automation to solve real-world problems and accelerate product performance."
)

# Save the document
output_path = "Tolulope_Adeleye_Resume_Updated.docx"
doc.save(output_path)

output_path
